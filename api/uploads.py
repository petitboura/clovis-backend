"""
Upload d'images (image de vitrine d'un agent, avatar de profil). Ajouté
le 2026-07-12 suite à un bug remonté par Bourama : les champs "URL image"
demandaient de coller un lien à la main, pas utilisable pour quelqu'un de
non-technique. Remplacés côté
frontend par un vrai bouton d'upload (components/ChampImage.tsx), qui
passe par ce endpoint.

L'upload passe TOUJOURS par ici (service role key), jamais directement du
navigateur vers Supabase Storage : pas de policy RLS sur storage.objects,
cohérent avec le reste du projet (aucune table n'a de policy non plus,
tout passe par le service role côté Python — voir la note dans la
migration Supabase `pivot_social_etape_b_tables`).
"""

import logging
import uuid
import base64
import os
import sys
import subprocess
import tempfile

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from google import genai

from api.auth import supabase, utilisateur_courant, get_secret

sys.path.append(os.path.join(os.path.dirname(__file__), "..", "core"))
from bibliotheque_fichiers import enregistrer_fichier, indexer_fichier_existant  # noqa: E402
from conversion_pdf import conversion_disponible, convertir_en_pdf  # noqa: E402
from core.erreurs import erreur_api

router = APIRouter(prefix="/api/uploads", tags=["uploads"])

BUCKET = "images-publiques"

TYPES_AUTORISES = {
    "image/jpeg": "jpg",
    "image/png": "png",
    "image/webp": "webp",
}

TAILLE_MAX_OCTETS = 5 * 1024 * 1024  # 5 Mo


@router.post("/image")
async def uploader_image(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    Upload une image (jpeg/png/webp, 5 Mo max) dans le bucket public
    `images-publiques`, sous le chemin `{user_id}/{uuid}.{extension}` —
    un dossier par utilisateur, pas de collision possible entre deux
    personnes qui uploadent au même moment. Renvoie l'URL publique,
    directement utilisable comme `image_vitrine_url` ou `avatar_url`.
    """
    if fichier.content_type not in TYPES_AUTORISES:
        raise erreur_api(400, "FORMAT_NON_SUPPORTE_JPEG_PNG_OU")

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_OCTETS:
        raise erreur_api(400, "IMAGE_TROP_LOURDE_5_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_VIDE")

    extension = TYPES_AUTORISES[fichier.content_type]
    chemin = f"{utilisateur.id}/{uuid.uuid4()}.{extension}"

    try:
        supabase.storage.from_(BUCKET).upload(
            chemin,
            contenu,
            {"content-type": fichier.content_type},
        )
    except Exception as e:
        logging.error(f"ERREUR SUPABASE STORAGE (upload {chemin}) : {e}")
        raise erreur_api(500, "ECHEC_DE_L_UPLOAD_REESSAIE")

    url = supabase.storage.from_(BUCKET).get_public_url(chemin)
    return {"url": url}


@router.post("/image-chat")
async def uploader_image_chat(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    Upload une image jointe à un message de chat (voir
    components/chat/BarreDeSaisie.tsx côté frontend, et
    core/main.py:chat() côté traitement -- routée vers Gemini, seul
    modèle multimodal de la cascade).

    Réutilise le même bucket `images-publiques` que uploader_image
    ci-dessus, sous un préfixe `chat/` dédié : ce projet n'a qu'un seul
    bucket Storage et aucune policy RLS dessus (tout passe par le service
    role, voir la note en tête de ce fichier) -- créer un bucket privé
    séparé est une étape d'infra Supabase à part, pas un changement de
    code. À faire si la confidentialité des images de conversation devient
    un sujet (le bucket actuel est public par construction : quiconque a
    l'URL peut voir l'image, comme pour un avatar ou une image de vitrine).

    Mêmes contraintes que /image (jpeg/png/webp, 5 Mo max) : Gemini
    n'accepte de toute façon que ces formats courants pour la vision.
    """
    if fichier.content_type not in TYPES_AUTORISES:
        raise erreur_api(400, "FORMAT_NON_SUPPORTE_JPEG_PNG_OU")

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_OCTETS:
        raise erreur_api(400, "IMAGE_TROP_LOURDE_5_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_VIDE")

    extension = TYPES_AUTORISES[fichier.content_type]
    chemin = f"chat/{utilisateur.id}/{uuid.uuid4()}.{extension}"

    try:
        supabase.storage.from_(BUCKET).upload(
            chemin,
            contenu,
            {"content-type": fichier.content_type},
        )
    except Exception as e:
        logging.error(f"ERREUR SUPABASE STORAGE (upload chat {chemin}) : {e}")
        raise erreur_api(500, "ECHEC_DE_L_UPLOAD_REESSAIE")

    url = supabase.storage.from_(BUCKET).get_public_url(chemin)

    # Persistance bibliothèque (2026-07-22, demande de Bourama : un
    # fichier uploadé par un utilisateur en chat ne doit plus être
    # "utilisé une fois puis perdu" -- l'IA doit pouvoir le retrouver
    # plus tard via chercher_fichier). Best-effort : un souci d'indexation
    # ne doit jamais empêcher l'upload lui-même de réussir (Gemini a déjà
    # ce dont il a besoin, l'url, à ce stade).
    try:
        indexer_fichier_existant(
            url_publique=url,
            chemin_stockage=chemin,
            nom_fichier=fichier.filename or "image.jpg",
            type_mime=fichier.content_type,
            niveau="utilisateur",
            uploade_par=utilisateur.id,
            user_id=utilisateur.id,
            description="Image envoyée en conversation",
            origine="chat",
            taille_octets=len(contenu),
        )
    except Exception as e:
        logging.warning(f"Indexation bibliothèque échouée pour image chat {chemin} (upload OK quand même) : {e}")

    return {"url": url}


PROMPT_EXTRACTION_FORMULE = (
    "Tu es un outil d'extraction de formule mathématique à partir d'une "
    "image (photo de feuille, manuel, écran, ou formule manuscrite). "
    "Renvoie UNIQUEMENT le LaTeX de la ou des formules visibles sur "
    "l'image, sans aucun texte d'accompagnement, sans délimiteurs $ ou "
    "$$, sans balises markdown ni bloc de code. S'il y a plusieurs "
    "formules distinctes, sépare-les par un retour à la ligne. Si "
    "l'image ne contient aucune formule mathématique lisible, réponds "
    "exactement : AUCUNE_FORMULE_DETECTEE"
)


@router.post("/extraire-formule")
async def extraire_formule(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    OCR ciblé formule (2026-07-26, demande de Bourama : priorité maths --
    "l'image convertie que tu peux éditer avant d'envoyer"). Contrairement
    à /image-chat, cette image n'est PAS destinée à rejoindre la
    conversation : elle sert uniquement à en extraire le LaTeX, renvoyé
    au frontend pour être ouvert dans EditeurFormule.tsx (éditable avant
    insertion). Pas de stockage Supabase ici -- l'image ne sert qu'à cet
    appel Gemini ponctuel, rien à retrouver plus tard dans la
    bibliothèque de fichiers (contrairement à une image envoyée en
    conversation).

    Mêmes contraintes de format/taille que /image-chat : Gemini
    (google-genai), seul modèle multimodal du projet, n'accepte de toute
    façon que jpeg/png/webp pour la vision.
    """
    if fichier.content_type not in TYPES_AUTORISES:
        raise erreur_api(400, "FORMAT_NON_SUPPORTE_JPEG_PNG_OU")

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_OCTETS:
        raise erreur_api(400, "IMAGE_TROP_LOURDE_5_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_VIDE")

    try:
        client_google = genai.Client(api_key=get_secret("GOOGLE_API_KEY"))
        reponse = client_google.models.generate_content(
            model="gemini-2.5-flash",
            contents=[{
                "role": "user",
                "parts": [
                    {"text": PROMPT_EXTRACTION_FORMULE},
                    {"inline_data": {"mime_type": fichier.content_type, "data": base64.b64encode(contenu).decode("utf-8")}},
                ],
            }],
        )
    except Exception as e:
        logging.error(f"ERREUR GEMINI (extraire-formule) : {e}")
        raise erreur_api(500, "ECHEC_DE_L_EXTRACTION_REESSAIE")

    latex = (reponse.text or "").strip().strip("`").strip()
    if not latex or latex == "AUCUNE_FORMULE_DETECTEE":
        raise erreur_api(422, "AUCUNE_FORMULE_DETECTEE_DANS_CETTE_IMAGE")

    return {"latex": latex}


# --- Documents (PDF/Word/Excel) : extraction texte, pas de stockage --------
#
# Contrairement aux images (qui doivent transiter par Gemini, donc par une
# URL publique), un document texte n'a pas besoin d'être stocké : on
# extrait son contenu ici et le frontend l'injecte tel quel dans le
# message texte envoyé à /api/chat (voir ChatIA.tsx:envoyerMessage).
# core/main.py n'a donc AUCUN changement à faire pour les documents -- le
# cascade Groq habituel les traite comme du texte normal.

TYPES_DOCUMENTS_AUTORISES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
}
TAILLE_MAX_DOCUMENT_OCTETS = 15 * 1024 * 1024  # 15 Mo
LONGUEUR_MAX_TEXTE_EXTRAIT = 30_000  # caractères, pour ne pas saturer le prompt système


def _extraire_texte_pdf(contenu_bytes):
    import io
    import PyPDF2

    reader = PyPDF2.PdfReader(io.BytesIO(contenu_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extraire_texte_docx(contenu_bytes):
    import io
    import docx

    document = docx.Document(io.BytesIO(contenu_bytes))
    morceaux = [p.text for p in document.paragraphs]

    # BUG corrigé le 2026-07-21 : les tableaux Word n'étaient jamais lus
    # (seuls document.paragraphs l'étaient) -- confirmé en test réel avec
    # un tableau de scores, dont aucune valeur ne remontait au modèle.
    # python-docx n'inclut pas les cellules de tableau dans .paragraphs,
    # il faut parcourir document.tables séparément.
    for table in document.tables:
        for ligne in table.rows:
            morceaux.append("\t".join(cellule.text for cellule in ligne.cells))

    return "\n".join(morceaux)


def _extraire_texte_xlsx(contenu_bytes):
    import io
    import openpyxl

    classeur = openpyxl.load_workbook(io.BytesIO(contenu_bytes), data_only=True)
    morceaux = []
    for feuille in classeur.worksheets:
        morceaux.append(f"--- Feuille : {feuille.title} ---")
        for ligne in feuille.iter_rows(values_only=True):
            morceaux.append(
                "\t".join("" if v is None else str(v) for v in ligne)
            )
    return "\n".join(morceaux)


@router.post("/document-chat")
async def uploader_document_chat(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    Extrait le texte d'un PDF/Word/Excel joint à un message de chat, pour
    injection dans le message avant envoi à /api/chat. Le fichier original
    est aussi stocké (voir enregistrer_fichier ci-dessous, depuis le
    2026-07-22) et, pour Word/Excel, converti en PDF pour permettre un
    aperçu visuel côté frontend (FichierChip.tsx) -- un PDF n'a pas besoin
    de conversion, il est déjà son propre aperçu.
    """
    if fichier.content_type not in TYPES_DOCUMENTS_AUTORISES:
        raise erreur_api(400, "FORMAT_NON_SUPPORTE_PDF_WORD_DOCX")

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_DOCUMENT_OCTETS:
        raise erreur_api(400, "DOCUMENT_TROP_LOURD_15_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_VIDE")

    extension = TYPES_DOCUMENTS_AUTORISES[fichier.content_type]
    try:
        if extension == "pdf":
            texte = _extraire_texte_pdf(contenu)
        elif extension == "docx":
            texte = _extraire_texte_docx(contenu)
        else:
            texte = _extraire_texte_xlsx(contenu)
    except Exception as e:
        logging.error(f"ERREUR EXTRACTION DOCUMENT ({fichier.filename}) : {e}")
        raise erreur_api(500, "ECHEC_DE_LA_LECTURE_DU_DOCUMENT")

    texte = texte.strip()
    if not texte:
        raise erreur_api(400, "AUCUN_TEXTE_TROUVE_DOCUMENT_SCANNE_IMAGE")

    tronque = len(texte) > LONGUEUR_MAX_TEXTE_EXTRAIT
    if tronque:
        texte = texte[:LONGUEUR_MAX_TEXTE_EXTRAIT]

    # Persistance bibliothèque (2026-07-22) : contrairement à avant, le
    # document original est maintenant gardé (pas seulement son texte
    # extrait), pour que l'IA puisse le retrouver et le redonner plus
    # tard via chercher_fichier. Best-effort : un souci ici ne doit
    # jamais faire échouer la réponse (le texte extrait reste le besoin
    # principal de cet endpoint).
    #
    # REVERT du 01/08 (Bourama : "je t'ai dit que les PDF uploadés dans
    # le chat n'y sont pas -- supprime-les de la page, eux restent que
    # dans le chat, tel qu'elle") : la tentative de vectorisation
    # (indexer_texte_bibliotheque) ajoutée plus tôt le même jour est
    # retirée. Un document envoyé en chat reste UNIQUEMENT un artefact de
    # cette conversation -- stocké et retrouvable par nom via
    # chercher_fichier comme avant, mais PAS injecté dans
    # documents_bibliotheque, et PAS listé sur la page Mon espace >
    # Bibliothèque (voir le filtre ajouté dans
    # api/bibliotheque_utilisateur.py:lister). Seul ce qui est ajouté à
    # la main depuis Mon espace fait partie de la base que
    # consulter_bibliotheque interroge.
    try:
        ligne = enregistrer_fichier(
            contenu=contenu,
            nom_fichier=fichier.filename or f"document.{extension}",
            type_mime=fichier.content_type,
            niveau="utilisateur",
            uploade_par=utilisateur.id,
            user_id=utilisateur.id,
            description="Document envoyé en conversation",
            origine="chat",
        )
        url_document = ligne["url_publique"]
    except Exception as e:
        logging.warning(f"Indexation bibliothèque échouée pour document chat {fichier.filename} (extraction OK quand même) : {e}")
        url_document = None

    # Aperçu PDF (25/07) : uniquement pour Word/Excel -- un PDF uploadé
    # est déjà consultable tel quel, pas besoin de conversion. Best-effort
    # comme le reste de cette fonction : CLOUDCONVERT_API_KEY absente ou
    # conversion échouée -> url_apercu reste None, le texte extrait et le
    # fichier original restent utilisables sans aperçu visuel.
    url_apercu = None
    if extension != "pdf" and conversion_disponible():
        try:
            pdf_bytes = convertir_en_pdf(contenu, fichier.filename or f"document.{extension}")
            chemin_apercu = f"apercus/{uuid.uuid4()}.pdf"
            supabase.storage.from_("images-publiques").upload(
                chemin_apercu, pdf_bytes, {"content-type": "application/pdf"}
            )
            url_apercu = supabase.storage.from_("images-publiques").get_public_url(chemin_apercu)
        except Exception as e:
            logging.warning(f"Aperçu PDF échoué pour document chat {fichier.filename} (extraction/stockage OK quand même) : {e}")

    return {"texte": texte, "tronque": tronque, "url": url_document, "url_apercu": url_apercu}


# --- Audio (dictée vocale) : transcription, pas de stockage ----------------
#
# Même logique que les documents : on transcrit puis on renvoie du texte
# brut, réinjecté dans le flux texte normal. Aucun changement requis dans
# core/main.py -- Whisper (via Groq, déjà le fournisseur du cascade
# principal) fait la transcription en amont, hors du cascade chat().

TAILLE_MAX_AUDIO_OCTETS = 20 * 1024 * 1024  # 20 Mo (limite Groq Whisper)

# Whisper hallucine des phrases précises quand l'audio reçu est silencieux
# ou quasi-vide, au lieu de renvoyer un texte vide -- artefact documenté de
# la communauté (hérité de ses données d'entraînement : sous-titres de
# vidéos, génériques de fin...). Confirmé le 2026-07-20 : "Sous-titrage
# Société Radio-Canada" est ressorti de façon répétée en dictée réelle.
# Sans ce filtre, ce texte halluciné serait envoyé tel quel comme si
# l'étudiant l'avait vraiment dit -- on le traite plutôt comme une
# transcription vide (même message d'erreur que "rien entendu").
PHRASES_HALLUCINEES_WHISPER = {
    "sous-titrage société radio-canada",
    "sous-titrage societe radio-canada",
    "sous-titres réalisés par la communauté d'amara.org",
    "sous-titres realises par la communaute d'amara.org",
    "merci d'avoir regardé cette vidéo",
    "merci d'avoir regardé la vidéo",
    "abonnez-vous à la chaîne",
    "www.tvsubtitles.net",
    "merci.",
    "sous-titres",
}


def _transcription_vraisemblable(texte):
    """
    True si le texte transcrit semble être une vraie parole captée, False
    s'il correspond à une hallucination Whisper connue sur audio silencieux
    (voir PHRASES_HALLUCINEES_WHISPER ci-dessus). Comparaison insensible à
    la casse/ponctuation de fin, pas de correspondance floue -- un faux
    négatif (vraie phrase qui ressemble à une hallucination) est jugé moins
    grave qu'un faux positif (hallucination envoyée comme si l'étudiant
    l'avait dite).
    """
    nettoye = texte.strip().lower().rstrip(".")
    return nettoye not in PHRASES_HALLUCINEES_WHISPER


@router.post("/audio-chat")
async def uploader_audio_chat(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    Transcrit un audio avec whisper-large-v3 via Groq (même fournisseur
    que le cascade texte, pas de nouvelle clé API à gérer). Sert deux
    cas : la dictée vocale (BarreDeSaisie.tsx, MediaRecorder côté
    navigateur, toujours actif) et, préparé le 2026-07-22 pour un vrai
    upload de fichier audio (mp3/wav/ogg/m4a) depuis le trombone -- pas
    de vérification de content-type ici (un blob MediaRecorder n'a pas
    toujours un type MIME standard), Whisper gère le format lui-même ou
    échoue proprement.
    """
    from groq import Groq

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_AUDIO_OCTETS:
        raise erreur_api(400, "AUDIO_TROP_LONG_20_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_AUDIO_VIDE")

    try:
        client_groq = Groq(api_key=get_secret("GROQ_API_KEY"))
        transcription = client_groq.audio.transcriptions.create(
            file=(fichier.filename or "audio.webm", contenu),
            model="whisper-large-v3",
            language="fr",
        )
    except Exception as e:
        logging.error(f"ERREUR TRANSCRIPTION AUDIO (Groq Whisper) : {e}")
        raise erreur_api(500, "ECHEC_DE_LA_TRANSCRIPTION_REESSAIE")

    texte = (transcription.text or "").strip()
    if not texte or not _transcription_vraisemblable(texte):
        raise erreur_api(400, "RIEN_N_A_ETE_COMPRIS_REESSAIE")

    # Persistance bibliothèque (2026-07-22) : même logique que
    # image/document/vidéo. S'applique aussi bien à la dictée micro qu'à
    # un futur vrai fichier audio uploadé -- best-effort, ne bloque jamais
    # la réponse (le texte transcrit est déjà prêt).
    try:
        ligne = enregistrer_fichier(
            contenu=contenu,
            nom_fichier=fichier.filename or "audio.webm",
            type_mime=fichier.content_type or "audio/webm",
            niveau="utilisateur",
            uploade_par=utilisateur.id,
            user_id=utilisateur.id,
            description="Audio envoyé en conversation",
            origine="chat",
        )
        url_audio = ligne["url_publique"]
    except Exception as e:
        logging.warning(f"Indexation bibliothèque échouée pour audio chat {fichier.filename} (transcription OK quand même) : {e}")
        url_audio = None

    return {"texte": texte, "url": url_audio}

    return {"texte": texte}


# --- Vidéo : extraction audio (transcription) + frames (analyse visuelle) -
#
# Contrairement au transcript YouTube (core/main.py:_lire_url, qui lit des
# sous-titres déjà écrits), ici on traite un vrai fichier vidéo uploadé :
# on en tire (1) une transcription audio via Whisper, comme pour la
# dictée vocale, et (2) quelques frames image envoyées à Gemini, comme
# pour une image simple. Rien n'est stocké -- fichiers temporaires,
# supprimés dans le `finally` de uploader_video_chat.
#
# LIMITE INFRA CONNUE : nécessite ffmpeg/ffprobe sur la machine qui
# exécute ce backend. Railway (Railpack) ne les installe PAS par défaut
# pour un projet Python -- voir railpack.json à la racine du repo pour
# les déclarer explicitement. Sans ce fichier, cet endpoint échoue en
# prod avec FileNotFoundError au premier subprocess.

TYPES_VIDEO_AUTORISES = {
    "video/mp4": "mp4",
    "video/webm": "webm",
    "video/quicktime": "mov",
}
TAILLE_MAX_VIDEO_OCTETS = 40 * 1024 * 1024  # 40 Mo
DUREE_MAX_VIDEO_SECONDES = 120  # 2 minutes -- garde ffmpeg + Whisper rapides
NB_FRAMES_VIDEO = 5


def _duree_video(chemin_video):
    resultat = subprocess.run(
        [
            "ffprobe", "-v", "error", "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1", chemin_video,
        ],
        capture_output=True, text=True, timeout=15, check=True,
    )
    return float(resultat.stdout.strip())


def _extraire_audio_video(chemin_video, chemin_audio_sortie):
    subprocess.run(
        [
            "ffmpeg", "-y", "-i", chemin_video, "-vn", "-ac", "1", "-ar", "16000",
            "-f", "wav", chemin_audio_sortie,
        ],
        capture_output=True, timeout=60, check=True,
    )


def _extraire_frames_video(chemin_video, duree, nb_frames=NB_FRAMES_VIDEO):
    """
    Extrait `nb_frames` images réparties uniformément sur la durée de la
    vidéo (jamais la toute première/dernière milliseconde, souvent noire
    ou floue) -- redimensionnées à 512px de large pour rester légères une
    fois envoyées à Gemini en base64. Une frame individuelle qui échoue
    (rare, décodage ponctuel) est simplement ignorée plutôt que de faire
    échouer toute la vidéo.
    """
    frames = []
    with tempfile.TemporaryDirectory() as dossier:
        for i in range(nb_frames):
            instant = duree * (i + 1) / (nb_frames + 1)
            chemin_frame = os.path.join(dossier, f"frame_{i}.jpg")
            try:
                subprocess.run(
                    [
                        "ffmpeg", "-y", "-ss", str(instant), "-i", chemin_video,
                        "-frames:v", "1", "-vf", "scale=512:-1", chemin_frame,
                    ],
                    capture_output=True, timeout=20, check=True,
                )
            except Exception as e:
                logging.error(f"ERREUR EXTRACTION FRAME {i} (t={instant:.1f}s) : {e}")
                continue
            if os.path.exists(chemin_frame):
                with open(chemin_frame, "rb") as f:
                    frames.append(f.read())
    return frames


@router.post("/video-chat")
async def uploader_video_chat(
    fichier: UploadFile = File(...),
    utilisateur=Depends(utilisateur_courant),
):
    """
    Traite une vidéo jointe à un message de chat : transcription audio
    (Whisper/Groq) + frames image (analysées ensuite par Gemini, voir
    core/main.py:chat(), paramètre images_base64). Renvoie
    {"transcript": str, "frames_base64": [str, ...]} -- le frontend injecte
    le transcript dans le texte du message et passe frames_base64 tel quel
    à /api/chat.
    """
    if fichier.content_type not in TYPES_VIDEO_AUTORISES:
        raise erreur_api(400, "FORMAT_NON_SUPPORTE_MP4_WEBM_OU")

    contenu = await fichier.read()
    if len(contenu) > TAILLE_MAX_VIDEO_OCTETS:
        raise erreur_api(400, "VIDEO_TROP_LOURDE_40_MO_MAX")
    if len(contenu) == 0:
        raise erreur_api(400, "FICHIER_VIDE")

    extension = TYPES_VIDEO_AUTORISES[fichier.content_type]
    with tempfile.NamedTemporaryFile(suffix=f".{extension}", delete=False) as f_video:
        f_video.write(contenu)
        chemin_video = f_video.name

    chemin_audio = None
    try:
        try:
            duree = _duree_video(chemin_video)
        except Exception as e:
            logging.error(f"ERREUR FFPROBE (durée vidéo) : {e}")
            raise erreur_api(500, "VIDEO_ILLISIBLE_REESSAIE_AVEC_UN_AUTRE")

        if duree > DUREE_MAX_VIDEO_SECONDES:
            raise erreur_api(400, "VIDEO_TROP_LONGUE", duree=int(duree), maximum=DUREE_MAX_VIDEO_SECONDES)

        transcript = ""
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f_audio:
            chemin_audio = f_audio.name
        try:
            _extraire_audio_video(chemin_video, chemin_audio)
            with open(chemin_audio, "rb") as f:
                from groq import Groq

                client_groq = Groq(api_key=get_secret("GROQ_API_KEY"))
                transcription = client_groq.audio.transcriptions.create(
                    file=("audio.wav", f.read()),
                    model="whisper-large-v3",
                    language="fr",
                )
                transcript = (transcription.text or "").strip()
                # Même filtre que uploader_audio_chat -- une vidéo sans son
                # exploitable (muette, piste corrompue) ne doit pas non
                # plus produire une hallucination Whisper faisant croire
                # qu'il y a un vrai contenu audio.
                if transcript and not _transcription_vraisemblable(transcript):
                    transcript = ""
        except Exception as e:
            # Pas bloquant : une vidéo sans son exploitable (muette, piste
            # audio corrompue) continue avec les frames seules.
            logging.error(f"ERREUR EXTRACTION/TRANSCRIPTION AUDIO VIDEO : {e}")

        try:
            frames = _extraire_frames_video(chemin_video, duree)
        except Exception as e:
            logging.error(f"ERREUR EXTRACTION FRAMES VIDEO : {e}")
            frames = []

        if not transcript and not frames:
            raise erreur_api(500, "IMPOSSIBLE_D_ANALYSER_CETTE_VIDEO_REESSAIE")

        frames_base64 = [base64.b64encode(f).decode("utf-8") for f in frames]

        # Persistance bibliothèque (2026-07-22) : même correctif que pour
        # les images/documents -- la vidéo originale était traitée puis
        # jetée, jamais retrouvable ensuite. Best-effort : ne doit jamais
        # faire échouer la réponse (transcript/frames sont déjà prêts).
        try:
            ligne = enregistrer_fichier(
                contenu=contenu,
                nom_fichier=fichier.filename or f"video.{extension}",
                type_mime=fichier.content_type,
                niveau="utilisateur",
                uploade_par=utilisateur.id,
                user_id=utilisateur.id,
                description="Vidéo envoyée en conversation",
                origine="chat",
            )
            url_video = ligne["url_publique"]
        except Exception as e:
            logging.warning(f"Indexation bibliothèque échouée pour vidéo chat {fichier.filename} (analyse OK quand même) : {e}")
            url_video = None

        return {"transcript": transcript, "frames_base64": frames_base64, "url": url_video}
    finally:
        if os.path.exists(chemin_video):
            os.remove(chemin_video)
        if chemin_audio and os.path.exists(chemin_audio):
            os.remove(chemin_audio)

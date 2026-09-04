"""
Cree le 04/09/2026, Bourama : vectorisation automatique en arriere-plan
de tout le contenu (hormis video) d'un dossier designe sur le telephone
(dossiers_designes_mobile.py), transfere via api/dossiers_designes.py.

Module DEDIE, distinct de core/file_attente_vectorisation.py (bibliotheque
perso/publique) -- meme inspiration (file d'attente, statuts, robustesse
au redemarrage) mais PAS le meme module, pour ne rien risquer sur le
systeme existant deja en production, et parce que la politique de reessai
est volontairement differente ici (voir plus bas).

Reutilise les fonctions d'extraction/description DEJA existantes
ailleurs dans le projet (aucune logique dupliquee pour rien) :
- PDF : meme logique que core/bibliotheque_rag.py::extraire_pages_pdf,
  dupliquee ici sur des bytes (voir extraire_pages_pdf_bytes plus bas --
  meme convention de duplication volontaire qu'ailleurs dans le projet
  pour ne pas creer de dependance croisee entre circuits)
- Image : core/description_multimedia.py::decrire_image_bibliotheque
- Audio : core/description_multimedia.py::transcrire_audio_bibliotheque
- Word (.docx) / Excel (.xlsx) : memes fonctions que api/uploads.py
  (_extraire_texte_docx/_extraire_texte_xlsx), dupliquees ici sur des
  bytes -- meme convention de duplication volontaire deja assumee entre
  api/uploads.py et core/lecture_fichier_mobile.py pour ce meme type de
  contenu (voir leurs docstrings respectifs).
- Texte brut (txt/md/csv/json/code...) : lu tel quel, meme liste
  d'extensions que core/lecture_fichier_mobile.py::EXTENSIONS_TEXTE_BRUT.
Le decoupage + embedding lui-meme passe par indexer_texte_bibliotheque-
like, mais ecrit dans documents_dossier_designe (table dediee) --
volontairement pas indexer_texte_bibliotheque (qui ecrit dans
documents_bibliotheque, scope different).

Politique de reessai (04/09, precisee par Bourama) : PAS de reessai
rapproche -- un echec passe DIRECTEMENT en statut "echec" des la
premiere tentative ratee (pas de boucle de retentatives immediates dans
le meme passage, contrairement a file_attente_vectorisation.py). Le seul
reessai est AUTOMATIQUE A FROID (apres COOLDOWN_REESSAI), et SANS
PLAFOND -- contrairement a MAX_TENTATIVES_AUTO de la bibliotheque perso,
un fichier repart indefiniment tant qu'il echoue, jusqu'a reussir. Pas de
bouton "reessayer" manuel pour l'instant (pas demande ici, a ajouter si
besoin).
"""

import io
import logging
import os
import sys
import tempfile
from datetime import datetime, timedelta, timezone

from supabase import create_client

sys.path.append(os.path.dirname(__file__))
from embeddings import vectoriser, decouper_texte  # noqa: E402
from description_multimedia import decrire_image_bibliotheque, transcrire_audio_bibliotheque  # noqa: E402

BUCKET_DOSSIERS_DESIGNES = "bibliotheque"  # meme bucket que la bibliotheque perso, sous-dossier "dossiers_designes/" (voir api/dossiers_designes.py)
TAILLE_LOT = 5  # meme garde-fou que file_attente_vectorisation.py -- un passage ne tourne jamais indefiniment
TAILLE_MAX_CHUNKS_PAR_DOCUMENT = 400  # meme plafond que bibliotheque_rag.py

# Reessai automatique a froid, SANS plafond de tentatives (voir docstring
# du module) -- seul un delai separe deux tentatives successives.
COOLDOWN_REESSAI = timedelta(minutes=15)

EXTENSIONS_TEXTE_BRUT = {
    "txt", "md", "csv", "json", "py", "js", "ts", "tsx", "jsx", "html", "css",
    "xml", "yaml", "yml", "java", "c", "cpp", "kt", "sh", "log",
}


def _get_secret(cle):
    return os.environ.get(cle)


supabase = create_client(_get_secret("SUPABASE_URL"), _get_secret("SUPABASE_SECRET"))


def _extension(nom_fichier: str) -> str:
    return nom_fichier.rsplit(".", 1)[-1].lower() if "." in (nom_fichier or "") else ""


def necessite_vectorisation(type_mime: str | None, nom_fichier: str) -> bool:
    """
    Tout est vectorise HORMIS la video (exclue explicitement par Bourama,
    trop couteux -- voir echange du 04/09). Le reste ("tout hormis
    video", demande explicite) : pdf, image, audio, docx, xlsx, texte
    brut. Un type totalement inconnu (ni mime reconnu ni extension texte
    brut) n'est pas vectorise -- pas d'erreur pour autant, juste "pret"
    directement (voir api/dossiers_designes.py).
    """
    if not type_mime:
        return _extension(nom_fichier) in EXTENSIONS_TEXTE_BRUT
    if type_mime.startswith("video/"):
        return False
    if type_mime == "application/pdf":
        return True
    if type_mime.startswith("image/"):
        return True
    if type_mime.startswith("audio/"):
        return True
    if type_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return True
    if type_mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return True
    if type_mime == "text/plain" or _extension(nom_fichier) in EXTENSIONS_TEXTE_BRUT:
        return True
    return False


def _extraire_texte_docx_bytes(contenu: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(contenu))
    morceaux = [p.text for p in document.paragraphs]
    for table in document.tables:
        for ligne in table.rows:
            morceaux.append("\t".join(cellule.text for cellule in ligne.cells))
    return "\n".join(morceaux)


def _extraire_texte_xlsx_bytes(contenu: bytes) -> str:
    import openpyxl

    classeur = openpyxl.load_workbook(io.BytesIO(contenu), data_only=True)
    morceaux = []
    for feuille in classeur.worksheets:
        morceaux.append(f"--- Feuille : {feuille.title} ---")
        for ligne in feuille.iter_rows(values_only=True):
            morceaux.append("\t".join("" if v is None else str(v) for v in ligne))
    return "\n".join(morceaux)


def _telecharger(chemin_stockage: str) -> bytes:
    return supabase.storage.from_(BUCKET_DOSSIERS_DESIGNES).download(chemin_stockage)


def _nettoyer_chunks_existants(fichier_id: str) -> None:
    """Meme principe que file_attente_vectorisation.py -- un traitement
    interrompu a moitie ne doit jamais laisser de chunks dupliques/incomplets."""
    supabase.table("documents_dossier_designe").delete().eq("fichier_id", fichier_id).execute()


def _indexer_texte(
    texte: str, fichier_id: str, user_id: str,
    page_debut=None, page_fin=None, timestamp_debut=None, timestamp_fin=None,
) -> int:
    morceaux = decouper_texte(texte)[:TAILLE_MAX_CHUNKS_PAR_DOCUMENT]
    lignes = []
    for morceau in morceaux:
        if not morceau.strip():
            continue
        embedding = vectoriser(morceau)
        lignes.append({
            "fichier_id": fichier_id,
            "user_id": user_id,
            "contenu": morceau,
            "embedding": embedding,
            "page_debut": page_debut,
            "page_fin": page_fin,
            "timestamp_debut": timestamp_debut,
            "timestamp_fin": timestamp_fin,
        })
    if lignes:
        supabase.table("documents_dossier_designe").insert(lignes).execute()
    return len(lignes)


def _vectoriser_fichier(ligne: dict) -> None:
    fichier_id = ligne["id"]
    user_id = ligne["user_id"]
    type_mime = ligne["type_mime"] or ""
    nom_fichier = ligne["nom_fichier"]
    extension = _extension(nom_fichier)
    contenu = _telecharger(ligne["chemin_stockage"])

    _nettoyer_chunks_existants(fichier_id)

    if type_mime == "application/pdf":
        for numero, texte_page in enumerate(extraire_pages_pdf_bytes(contenu), start=1):
            if texte_page.strip():
                _indexer_texte(texte_page, fichier_id, user_id, page_debut=numero, page_fin=numero)
    elif type_mime.startswith("image/"):
        description = decrire_image_bibliotheque(contenu, type_mime)
        if description:
            _indexer_texte(description, fichier_id, user_id)
    elif type_mime.startswith("audio/"):
        segments = transcrire_audio_bibliotheque(contenu, nom_fichier)
        for segment in segments or []:
            texte = (segment.get("text") or "").strip()
            if texte:
                _indexer_texte(texte, fichier_id, user_id, timestamp_debut=segment.get("start"), timestamp_fin=segment.get("end"))
    elif type_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        texte = _extraire_texte_docx_bytes(contenu)
        if texte.strip():
            _indexer_texte(texte, fichier_id, user_id)
    elif type_mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        texte = _extraire_texte_xlsx_bytes(contenu)
        if texte.strip():
            _indexer_texte(texte, fichier_id, user_id)
    elif type_mime == "text/plain" or extension in EXTENSIONS_TEXTE_BRUT:
        texte = contenu.decode("utf-8", errors="ignore")
        if texte.strip():
            _indexer_texte(texte, fichier_id, user_id)


def extraire_pages_pdf_bytes(contenu: bytes) -> list[str]:
    """Meme logique que bibliotheque_rag.py::extraire_pages_pdf, mais sur des bytes (pas un chemin de fichier -- le contenu vient de Supabase Storage, jamais ecrit sur disque)."""
    import PyPDF2

    pages = []
    reader = PyPDF2.PdfReader(io.BytesIO(contenu))
    for page in reader.pages:
        pages.append((page.extract_text() or "").replace("\x00", ""))
    return pages


def remettre_en_attente_bloques() -> None:
    """Appelee une fois au demarrage du process -- voir docstring de file_attente_vectorisation.py::remettre_en_attente_bloques (meme raison : Railway redeploie a chaque push)."""
    try:
        supabase.table("fichiers_dossier_designe").update({"statut_vectorisation": "en_attente"}).eq(
            "statut_vectorisation", "en_cours"
        ).execute()
    except Exception as e:
        logging.error(f"ERREUR remise en attente au demarrage (dossiers designes) : {e}")


COLONNES = "id, user_id, chemin_stockage, nom_fichier, type_mime, tentatives_vectorisation"


def traiter_file_attente_une_fois() -> int:
    """
    Traite jusqu'a TAILLE_LOT fichiers "en_attente", du plus ancien au
    plus recent. PAS de reessai rapproche (voir docstring du module) :
    un echec passe direct en "echec", relancer_echecs_a_froid ci-dessous
    s'en charge ensuite. Renvoie le nombre de fichiers traites.
    """
    try:
        lignes = (
            supabase.table("fichiers_dossier_designe")
            .select(COLONNES)
            .eq("statut_vectorisation", "en_attente")
            .order("created_at")
            .limit(TAILLE_LOT)
            .execute()
        ).data or []
    except Exception as e:
        logging.error(f"ERREUR lecture file d'attente (dossiers designes) : {e}")
        return 0

    for ligne in lignes:
        fichier_id = ligne["id"]
        maintenant_iso = datetime.now(timezone.utc).isoformat()
        try:
            supabase.table("fichiers_dossier_designe").update({"statut_vectorisation": "en_cours"}).eq("id", fichier_id).execute()
            _vectoriser_fichier(ligne)
            supabase.table("fichiers_dossier_designe").update({
                "statut_vectorisation": "pret",
                "erreur_vectorisation": None,
                "derniere_tentative_vectorisation_a": maintenant_iso,
            }).eq("id", fichier_id).execute()
        except Exception as e:
            tentatives = (ligne.get("tentatives_vectorisation") or 0) + 1
            logging.error(f"ERREUR vectorisation (dossiers designes, fichier_id={fichier_id}, tentative {tentatives}) : {e}")
            try:
                supabase.table("fichiers_dossier_designe").update({
                    "statut_vectorisation": "echec",
                    "tentatives_vectorisation": tentatives,
                    "erreur_vectorisation": str(e)[:500],
                    "derniere_tentative_vectorisation_a": maintenant_iso,
                }).eq("id", fichier_id).execute()
            except Exception as e2:
                logging.error(f"ERREUR mise a jour statut echec (dossiers designes, fichier_id={fichier_id}) : {e2}")

    return len(lignes)


def chercher_dossiers_designes(question: str, user_id: str, match_count: int = 5) -> list:
    """
    Recherche semantique dans TOUT le contenu deja vectorise des dossiers
    designes de `user_id` (tous dossiers confondus -- pas de filtre par
    dossier_nom ici, voir recherche_dossiers_designes en SQL). Meme
    principe que chercher_bibliotheque (core/bibliotheque_rag.py), sur la
    table dediee documents_dossier_designe.

    Renvoie une liste de {contenu, similarite, fichier_id, nom_fichier,
    dossier_nom, chemin, url_publique, type_mime, page_debut, page_fin,
    timestamp_debut, timestamp_fin} triee par pertinence. `chemin` est la
    liste ordonnee des sous-dossiers depuis la racine designee (jamais le
    nom du fichier), voir migrations/2026_09_04_dossiers_designes_
    vectorisation.sql.
    """
    if not user_id:
        logging.error("chercher_dossiers_designes appele sans user_id : renvoie vide.")
        return []

    try:
        vecteur = vectoriser(question, task_type="RETRIEVAL_QUERY")
    except Exception as e:
        logging.error(f"ERREUR VECTORISATION dossiers designes (Gemini) : {e}")
        return []

    try:
        return supabase.rpc(
            "recherche_dossiers_designes",
            {"query_embedding": vecteur, "match_count": match_count, "p_user_id": user_id},
        ).execute().data or []
    except Exception as e:
        logging.error(f"ERREUR SUPABASE RPC recherche_dossiers_designes (user_id={user_id}) : {e}")
        return []


def formater_source_dossier_designe(r: dict) -> str | None:
    """
    Meme role que formater_source_bibliotheque (core/bibliotheque_rag.py),
    adapte aux dossiers designes : ajoute le chemin (dossier_nom + sous-
    dossiers) pour que l'IA sache d'ou vient chaque extrait, en plus de la
    page/du timestamp quand ce chunk en a un.
    """
    if not (r.get("nom_fichier") and r.get("dossier_nom")):
        return None
    reperage = ""
    if r.get("page_debut") is not None:
        if r.get("page_fin") and r["page_fin"] != r["page_debut"]:
            reperage = f", page {r['page_debut']}-{r['page_fin']}"
        else:
            reperage = f", page {r['page_debut']}"
    elif r.get("timestamp_debut") is not None:
        debut = int(r["timestamp_debut"])
        reperage = f", à {debut // 60:02d}:{debut % 60:02d}"
    chemin = r.get("chemin") or []
    emplacement = " / ".join([r["dossier_nom"], *chemin]) if chemin else r["dossier_nom"]
    lien = f", {r['url_publique']}" if r.get("url_publique") else ""
    type_mime = r.get("type_mime") or ""
    return f"(Source : {r['nom_fichier']}{reperage}, dossier {emplacement}{lien}, {type_mime})"


def relancer_echecs_a_froid() -> int:
    """Reessai automatique a froid, SANS plafond (voir docstring du module) -- seul le cooldown separe deux tentatives."""
    seuil = (datetime.now(timezone.utc) - COOLDOWN_REESSAI).isoformat()
    try:
        resultat = (
            supabase.table("fichiers_dossier_designe")
            .update({"statut_vectorisation": "en_attente"})
            .eq("statut_vectorisation", "echec")
            .lt("derniere_tentative_vectorisation_a", seuil)
            .execute()
        )
        return len(resultat.data or [])
    except Exception as e:
        logging.error(f"ERREUR reessai automatique a froid (dossiers designes) : {e}")
        return 0

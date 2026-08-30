"""
Cree le 30/08/2026, Bourama : Lot 4 Partie 3 (app mobile), chantier
"Exploration de dossier en temps reel" (voir 00-commun-exploration-dossier.md
et 04-lecture-contenu.md a la racine du depot).

Lecture reelle du contenu d'un fichier trouve via core/exploration_dossier_mobile.py
(Lot 3) : le telephone envoie le contenu brut, encode en base64 dans la
reponse JSON du canal temps reel (core/canal_temps_reel.py, Lot 1). Ce
module choisit le traitement a appliquer selon le type MIME et renvoie un
texte exploitable par l'agent. Il ne contient aucune logique de canal
ou de dossier, uniquement la lecture.

Reutilise volontairement les circuits DEJA existants ailleurs dans le
projet plutot que d'en reconstruire :
- Image : core/description_multimedia.py::decrire_image_bibliotheque
  (vision Gemini), deja utilisee pour rendre les images de la
  bibliotheque cherchables.
- Audio : core/description_multimedia.py::transcrire_audio_bibliotheque
  (Whisper Groq), meme filtrage des hallucinations connues
  (PHRASES_HALLUCINEES_WHISPER) deja en place.
- PDF / Word (.docx) : PyPDF2 / python-docx, meme logique que
  core/bibliotheque_rag.py::extraire_texte_pdf et
  api/uploads.py::_extraire_texte_docx, dupliquee ici sur des bytes
  (et non un chemin de fichier, puisque le contenu arrive par le canal
  temps reel, jamais ecrit sur disque), meme convention de duplication
  volontaire deja assumee entre ces deux modules pour ne pas creer de
  dependance croisee entre circuits (voir leurs docstrings respectifs).
- Excel (.xlsx) : openpyxl, meme logique que
  api/uploads.py::_extraire_texte_xlsx. Pas nomme explicitement dans
  04-lecture-contenu.md mais deja disponible dans le projet et couvert
  par "tous les formats bureautiques" (00-commun-exploration-dossier.md,
  "types de fichiers a couvrir : tous, rien d'exclu a priori").

Point tranche avec Bourama le 30/08/2026 (voir 04-lecture-contenu.md,
"Point technique a trancher avec Bourama avant de coder ce lot") :
PAS de lecture pour un fichier trop volumineux pour l'instant : l'agent
le dit clairement a l'etudiant, capacite prevue plus tard, plutot que de
tenter une lecture qui risquerait de saturer le canal WebSocket. Seuils
PAR TYPE, alignes sur ceux deja en place ailleurs dans le projet pour le
meme type de fichier (api/uploads.py), pas de nouveau seuil invente ici,
sauf pour le texte brut ou aucun seuil n'existait deja.
"""

import base64
import logging

from core.description_multimedia import (
    decrire_image_bibliotheque,
    transcrire_audio_bibliotheque,
)

# Seuils alignes sur ceux deja en place dans api/uploads.py pour le meme
# type de fichier cote upload de chat. Volontairement pas de nouveau
# seuil invente pour ces types-la.
TAILLE_MAX_IMAGE_OCTETS = 5 * 1024 * 1024  # 5 Mo, meme limite que api/uploads.py (upload image chat)
TAILLE_MAX_DOCUMENT_OCTETS = 15 * 1024 * 1024  # 15 Mo, meme limite que api/uploads.py (PDF/Word/Excel chat)
TAILLE_MAX_AUDIO_OCTETS = 20 * 1024 * 1024  # 20 Mo, meme limite que api/uploads.py (limite Groq Whisper)
TAILLE_MAX_TEXTE_OCTETS = 2 * 1024 * 1024  # 2 Mo, aucun seuil existant ailleurs pour du texte brut, seuil prudent choisi ici

TYPES_IMAGE = {
    "image/jpeg", "image/jpg", "image/png", "image/webp", "image/heic", "image/heif",
}
TYPES_AUDIO = {
    "audio/mpeg", "audio/mp3", "audio/mp4", "audio/m4a", "audio/wav", "audio/x-wav",
    "audio/aac", "audio/ogg",
}
TYPE_PDF = "application/pdf"
TYPE_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TYPE_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# Extensions de texte brut courantes (cours, code, notes), lues telles
# quelles sans aucun traitement, comme demande dans 04-lecture-contenu.md.
EXTENSIONS_TEXTE_BRUT = {
    "txt", "md", "csv", "json", "py", "js", "ts", "tsx", "jsx", "html", "css",
    "xml", "yaml", "yml", "java", "c", "cpp", "kt", "sh", "log",
}


def _extension(nom_fichier: str) -> str:
    return nom_fichier.rsplit(".", 1)[-1].lower() if "." in (nom_fichier or "") else ""


def _extraire_texte_pdf_bytes(contenu: bytes) -> str:
    import io
    import PyPDF2

    reader = PyPDF2.PdfReader(io.BytesIO(contenu))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extraire_texte_docx_bytes(contenu: bytes) -> str:
    import io
    import docx

    document = docx.Document(io.BytesIO(contenu))
    morceaux = [p.text for p in document.paragraphs]

    # Meme correction que api/uploads.py::_extraire_texte_docx (21/07) :
    # python-docx n'inclut pas les cellules de tableau dans .paragraphs.
    for table in document.tables:
        for ligne in table.rows:
            morceaux.append("\t".join(cellule.text for cellule in ligne.cells))

    return "\n".join(morceaux)


def _extraire_texte_xlsx_bytes(contenu: bytes) -> str:
    import io
    import openpyxl

    classeur = openpyxl.load_workbook(io.BytesIO(contenu), data_only=True)
    morceaux = []
    for feuille in classeur.worksheets:
        morceaux.append(f"=== Feuille : {feuille.title} ===")
        for ligne in feuille.iter_rows(values_only=True):
            morceaux.append("\t".join("" if v is None else str(v) for v in ligne))
    return "\n".join(morceaux)


def fichier_trop_volumineux(type_mime: str, taille_octets: int | None) -> bool:
    """
    True si `taille_octets` depasse le seuil du type concerne. False si
    la taille est inconnue (on ne bloque jamais sur une info manquante,
    seulement sur une taille reellement mesuree trop grande : l'echec
    de lecture eventuel remontera alors naturellement via lire_contenu_fichier).
    """
    if taille_octets is None:
        return False
    if type_mime in TYPES_IMAGE:
        return taille_octets > TAILLE_MAX_IMAGE_OCTETS
    if type_mime in TYPES_AUDIO:
        return taille_octets > TAILLE_MAX_AUDIO_OCTETS
    if type_mime in (TYPE_PDF, TYPE_DOCX, TYPE_XLSX):
        return taille_octets > TAILLE_MAX_DOCUMENT_OCTETS
    return taille_octets > TAILLE_MAX_TEXTE_OCTETS


def lire_contenu_fichier(contenu_base64: str, type_mime: str, nom_fichier: str) -> dict:
    """
    Decode le contenu base64 recu du telephone et applique le traitement
    adapte au type MIME. Renvoie toujours un dict :
    - {"texte": "..."} en cas de succes (texte brut, texte extrait d'un
      PDF/Word/Excel, description d'image, ou transcription audio) ;
    - {"erreur": "..."} si le traitement echoue ou si le type n'est pas
      (encore) pris en charge.

    Un seul essai, jamais de reessai automatique en cas d'echec, meme
    convention que le reste du chantier (voir
    00-commun-exploration-dossier.md, "un seul essai par
    recherche/lecture").
    """
    try:
        contenu = base64.b64decode(contenu_base64)
    except Exception as e:
        logging.error(f"ERREUR decodage base64 lecture fichier ({nom_fichier}) : {e}")
        return {"erreur": "Contenu du fichier illisible (erreur de transfert)."}

    try:
        if type_mime in TYPES_IMAGE:
            description = decrire_image_bibliotheque(contenu, type_mime)
            if description is None:
                return {"erreur": "Impossible de décrire cette image."}
            return {"texte": description}

        if type_mime in TYPES_AUDIO:
            segments = transcrire_audio_bibliotheque(contenu, nom_fichier)
            if not segments:
                return {"erreur": "Impossible de transcrire cet audio (silencieux ou illisible)."}
            return {"texte": " ".join(segment["text"] for segment in segments)}

        if type_mime == TYPE_PDF:
            texte = _extraire_texte_pdf_bytes(contenu).strip()
            if not texte:
                return {"erreur": "Aucun texte trouvé dans ce PDF (probablement un scan sans OCR)."}
            return {"texte": texte}

        if type_mime == TYPE_DOCX:
            texte = _extraire_texte_docx_bytes(contenu).strip()
            if not texte:
                return {"erreur": "Ce document Word semble vide."}
            return {"texte": texte}

        if type_mime == TYPE_XLSX:
            texte = _extraire_texte_xlsx_bytes(contenu).strip()
            if not texte:
                return {"erreur": "Ce fichier Excel semble vide."}
            return {"texte": texte}

        if _extension(nom_fichier) in EXTENSIONS_TEXTE_BRUT or (type_mime or "").startswith("text/"):
            try:
                return {"texte": contenu.decode("utf-8")}
            except UnicodeDecodeError:
                return {"texte": contenu.decode("latin-1", errors="replace")}

        return {
            "erreur": f"Type de fichier non pris en charge pour la lecture pour l'instant ({type_mime or 'inconnu'})."
        }
    except Exception as e:
        logging.error(f"ERREUR lecture fichier ({nom_fichier}, {type_mime}) : {e}")
        return {"erreur": "Échec de la lecture de ce fichier."}
